#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中医病症-处方提取器
专门提取Word文档中的具体病症对应的处方内容
"""

import os
import re
import sys
import json
from docx import Document
from typing import Dict, List, Any, Optional, Tuple
import logging

# 添加项目路径
sys.path.append('/opt/tcm-ai')

logger = logging.getLogger(__name__)

class PrescriptionExtractor:
    """中医处方提取器"""
    
    def __init__(self):
        # 中医证型/病症模式
        self.syndrome_patterns = {
            # 感冒类
            'wind_cold': [
                '风寒感冒', '风寒外感', '风寒束表', '外感风寒',
                '恶寒发热', '无汗恶寒', '风寒表证'
            ],
            'wind_heat': [
                '风热感冒', '风热外感', '外感风热', '风热表证',
                '发热恶风', '咽红肿痛', '风热上扰'
            ],
            
            # 咳嗽类
            'lung_cold': [
                '风寒咳嗽', '肺寒咳嗽', '寒痰咳嗽', '外感风寒咳嗽'
            ],
            'lung_heat': [
                '风热咳嗽', '肺热咳嗽', '痰热咳嗽', '肺燥干咳'
            ],
            'phlegm_damp': [
                '痰湿咳嗽', '痰浊阻肺', '湿痰咳嗽'
            ],
            
            # 高血压类
            'liver_yang': [
                '肝阳上亢', '肝火上炎', '肝阳头痛', '肝风内动'
            ],
            'kidney_yin_def': [
                '肾阴虚', '肝肾阴虚', '阴虚阳亢', '水不涵木'
            ],
            'phlegm_blood_stasis': [
                '痰瘀互结', '痰湿中阻', '血瘀头痛'
            ],
            
            # 消化系统
            'spleen_stomach_weak': [
                '脾胃虚弱', '脾胃气虚', '中气不足', '脾气虚'
            ],
            'stomach_heat': [
                '胃热', '胃火', '胃燥', '胃阴虚'
            ],
            
            # 妇科
            'kidney_def': [
                '肾虚', '肾气虚', '肾阳虚', '肾精不足'
            ],
            'liver_qi_stag': [
                '肝气郁结', '气滞血瘀', '肝郁脾虚'
            ]
        }
        
        # 方剂名称模式
        self.prescription_name_patterns = [
            r'[\\u4e00-\\u9fff]{2,8}汤',
            r'[\\u4e00-\\u9fff]{2,8}散',
            r'[\\u4e00-\\u9fff]{2,8}丸',
            r'[\\u4e00-\\u9fff]{2,8}膏',
            r'[\\u4e00-\\u9fff]{2,8}饮',
            r'[\\u4e00-\\u9fff]{2,8}方'
        ]
        
        # 药物剂量模式
        self.herb_dose_patterns = [
            r'([\\u4e00-\\u9fff]{2,6})\\s*(\\d+(?:\\.\\d+)?)\\s*([克钱g])',
            r'([\\u4e00-\\u9fff]{2,6})\\s*(\\d+(?:\\.\\d+)?)\\s*([两斤])',
            r'([\\u4e00-\\u9fff]{2,6})\\s*(\\d+)\\s*([片枚粒])'
        ]
        
        # 处方结构关键词
        self.prescription_keywords = {
            'syndrome': ['证型', '辨证', '证候', '病机'],
            'treatment': ['治法', '治则', '治疗原则'],
            'prescription': ['方药', '处方', '方剂', '组成'],
            'usage': ['用法', '服法', '煎服法', '用量'],
            'modification': ['加减', '随证加减', '临证加减'],
            'indication': ['主治', '功效', '适用', '治疗']
        }
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """从docx提取文本"""
        try:
            doc = Document(docx_path)
            text_content = []
            
            # 提取段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
            
            # 提取表格（很多处方在表格中）
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\\n'.join(text_content)
        except Exception as e:
            logger.error(f"提取docx文本失败 {docx_path}: {e}")
            return ""
    
    def identify_syndromes(self, text: str) -> List[Dict[str, Any]]:
        """识别文档中的证型/病症"""
        syndromes_found = []
        
        for syndrome_type, patterns in self.syndrome_patterns.items():
            for pattern in patterns:
                # 查找证型及其上下文
                matches = []
                for match in re.finditer(pattern, text):
                    start = max(0, match.start() - 100)
                    end = min(len(text), match.end() + 200)
                    context = text[start:end]
                    
                    matches.append({
                        'syndrome_name': pattern,
                        'syndrome_type': syndrome_type,
                        'position': match.start(),
                        'context': context.replace('\\n', ' ')
                    })
                
                if matches:
                    syndromes_found.extend(matches)
        
        # 去重并按位置排序
        unique_syndromes = []
        seen_positions = set()
        
        for syndrome in sorted(syndromes_found, key=lambda x: x['position']):
            # 避免重复位置
            if not any(abs(syndrome['position'] - pos) < 50 for pos in seen_positions):
                unique_syndromes.append(syndrome)
                seen_positions.add(syndrome['position'])
        
        return unique_syndromes[:10]  # 最多返回10个
    
    def extract_prescriptions_for_syndrome(self, text: str, syndrome: Dict[str, Any]) -> Dict[str, Any]:
        """为特定证型提取处方信息"""
        prescription_info = {
            'syndrome': syndrome,
            'prescriptions': [],
            'treatments': [],
            'modifications': []
        }
        
        # 在证型附近查找处方信息
        syndrome_pos = syndrome['position']
        
        # 定义搜索范围（证型前后1000字符）
        search_start = max(0, syndrome_pos - 500)
        search_end = min(len(text), syndrome_pos + 1500)
        search_text = text[search_start:search_end]
        
        # 1. 查找治法
        treatments = self._extract_treatments(search_text)
        prescription_info['treatments'] = treatments
        
        # 2. 查找方剂名称
        prescription_names = self._extract_prescription_names(search_text)
        
        # 3. 查找药物组成
        herb_compositions = self._extract_herb_compositions(search_text)
        
        # 4. 查找用法用量
        usages = self._extract_usages(search_text)
        
        # 5. 查找加减方法
        modifications = self._extract_modifications(search_text)
        prescription_info['modifications'] = modifications
        
        # 组合处方信息
        for i, name in enumerate(prescription_names):
            prescription = {
                'name': name,
                'herbs': herb_compositions[i] if i < len(herb_compositions) else [],
                'usage': usages[i] if i < len(usages) else '',
                'position_in_text': search_start + search_text.find(name) if name in search_text else -1
            }
            prescription_info['prescriptions'].append(prescription)
        
        # 如果没找到明确的方剂名，但有药物组成，创建匿名处方
        if not prescription_names and herb_compositions:
            for i, herbs in enumerate(herb_compositions):
                prescription = {
                    'name': f"处方{i+1}",
                    'herbs': herbs,
                    'usage': usages[i] if i < len(usages) else '',
                    'position_in_text': -1
                }
                prescription_info['prescriptions'].append(prescription)
        
        return prescription_info
    
    def _extract_treatments(self, text: str) -> List[str]:
        """提取治法"""
        treatments = []
        
        # 查找治法模式
        for keyword in self.prescription_keywords['treatment']:
            pattern = f'{keyword}[：:]([^。\\n]{{5,50}})'
            matches = re.findall(pattern, text)
            treatments.extend(matches)
        
        # 查找常见治法词汇
        treatment_terms = [
            '疏风散寒', '清热解毒', '宣肺止咳', '化痰止咳', '润肺止咳',
            '平肝潜阳', '滋阴降火', '健脾益气', '温中散寒', '活血化瘀',
            '补气养血', '滋阴润燥', '温阳化气', '疏肝理气'
        ]
        
        for term in treatment_terms:
            if term in text:
                treatments.append(term)
        
        return list(set(treatments))  # 去重
    
    def _extract_prescription_names(self, text: str) -> List[str]:
        """提取方剂名称"""
        names = []
        
        for pattern in self.prescription_name_patterns:
            matches = re.findall(pattern, text)
            names.extend(matches)
        
        # 过滤常见的非方剂名词汇
        exclude_words = ['症状', '疾病', '患者', '医生', '治疗', '方法', '效果', '时间']
        filtered_names = [name for name in names if not any(word in name for word in exclude_words)]
        
        return list(set(filtered_names))  # 去重
    
    def _extract_herb_compositions(self, text: str) -> List[List[Dict[str, str]]]:
        """提取药物组成"""
        compositions = []
        
        # 查找药物组成段落
        composition_sections = []
        
        # 方法1：查找"组成"或"方药"后的内容
        for keyword in ['组成', '方药', '处方']:
            pattern = f'{keyword}[：:]([^。]{{20,300}})'
            matches = re.findall(pattern, text, re.DOTALL)
            composition_sections.extend(matches)
        
        # 方法2：查找连续的药物名称+剂量
        herb_sections = re.findall(r'([\\u4e00-\\u9fff\\s\\d克钱g,、，]{50,200})', text)
        composition_sections.extend(herb_sections)
        
        for section in composition_sections:
            herbs = self._parse_herb_list(section)
            if len(herbs) >= 3:  # 至少3味药才认为是完整处方
                compositions.append(herbs)
        
        return compositions[:3]  # 最多返回3个处方
    
    def _parse_herb_list(self, text: str) -> List[Dict[str, str]]:
        """解析药物列表"""
        herbs = []
        
        for pattern in self.herb_dose_patterns:
            matches = re.findall(pattern, text)
            for herb_name, dose, unit in matches:
                # 过滤明显不是药名的词
                if len(herb_name) >= 2 and not any(char.isdigit() for char in herb_name):
                    herbs.append({
                        'name': herb_name,
                        'dose': dose,
                        'unit': unit
                    })
        
        # 去重（同一药物只保留第一次出现）
        seen_herbs = set()
        unique_herbs = []
        for herb in herbs:
            if herb['name'] not in seen_herbs:
                unique_herbs.append(herb)
                seen_herbs.add(herb['name'])
        
        return unique_herbs
    
    def _extract_usages(self, text: str) -> List[str]:
        """提取用法用量"""
        usages = []
        
        usage_patterns = [
            r'用法[：:]([^。\\n]{10,100})',
            r'服法[：:]([^。\\n]{10,100})',
            r'每日[^。\\n]{5,50}',
            r'水煎服[^。\\n]{0,30}',
            r'[每]?日?[一二三]?[次剂][^。\\n]{5,30}'
        ]
        
        for pattern in usage_patterns:
            matches = re.findall(pattern, text)
            usages.extend(matches)
        
        return list(set(usages))
    
    def _extract_modifications(self, text: str) -> List[str]:
        """提取加减方法"""
        modifications = []
        
        modification_patterns = [
            r'加减[：:]([^。\\n]{10,100})',
            r'随证加减[：:]([^。\\n]{10,100})',
            r'若[^。\\n]{5,50}加[^。\\n]{5,30}',
            r'如[^。\\n]{5,50}加[^。\\n]{5,30}'
        ]
        
        for pattern in modification_patterns:
            matches = re.findall(pattern, text)
            modifications.extend(matches)
        
        return modifications
    
    def process_document(self, docx_path: str) -> Dict[str, Any]:
        """处理单个文档，提取所有病症-处方信息"""
        result = {
            'success': False,
            'filename': os.path.basename(docx_path),
            'syndromes_count': 0,
            'prescriptions_count': 0,
            'extracted_data': [],
            'error_message': ''
        }
        
        try:
            # 提取文本
            text = self.extract_text_from_docx(docx_path)
            if not text:
                result['error_message'] = "无法提取文档内容"
                return result
            
            # 识别证型
            syndromes = self.identify_syndromes(text)
            result['syndromes_count'] = len(syndromes)
            
            if not syndromes:
                result['error_message'] = "未识别到证型/病症"
                return result
            
            # 为每个证型提取处方
            total_prescriptions = 0
            for syndrome in syndromes:
                prescription_info = self.extract_prescriptions_for_syndrome(text, syndrome)
                if prescription_info['prescriptions']:
                    result['extracted_data'].append(prescription_info)
                    total_prescriptions += len(prescription_info['prescriptions'])
            
            result['prescriptions_count'] = total_prescriptions
            result['success'] = True
            
            logger.info(f"处理完成 {result['filename']}: {len(syndromes)}个证型, {total_prescriptions}个处方")
            
        except Exception as e:
            result['error_message'] = str(e)
            logger.error(f"处理文档失败 {docx_path}: {e}")
        
        return result

def main():
    """测试处方提取器"""
    extractor = PrescriptionExtractor()
    
    # 测试文件
    test_files = [
        '/opt/tcm-ai/all_tcm_docs/感冒.docx',
        '/opt/tcm-ai/all_tcm_docs/咳嗽.docx',
        '/opt/tcm-ai/all_tcm_docs/高血压.docx'
    ]
    
    print("🔬 中医处方提取器测试")
    print("=" * 50)
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print(f"\\n📄 处理: {test_file.split('/')[-1]}")
            result = extractor.process_document(test_file)
            
            if result['success']:
                print(f"✅ 成功: {result['syndromes_count']}个证型, {result['prescriptions_count']}个处方")
                
                # 显示前2个提取结果
                for i, data in enumerate(result['extracted_data'][:2]):
                    syndrome = data['syndrome']
                    print(f"\\n证型{i+1}: {syndrome['syndrome_name']}")
                    print(f"  治法: {', '.join(data['treatments'])}")
                    
                    for j, prescription in enumerate(data['prescriptions'][:1]):
                        print(f"  处方{j+1}: {prescription['name']}")
                        if prescription['herbs']:
                            herb_list = ', '.join([f"{h['name']}{h['dose']}{h['unit']}" for h in prescription['herbs'][:5]])
                            print(f"    药物: {herb_list}")
            else:
                print(f"❌ 失败: {result['error_message']}")

if __name__ == "__main__":
    main()