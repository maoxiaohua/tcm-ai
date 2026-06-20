#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中医文档预处理模块
处理docx文档，提取文本并进行中医专业清洗
"""

import os
import re
import jieba
from docx import Document
from typing import List, Dict, Any, Optional, Tuple
import logging

logger = logging.getLogger(__name__)

class TCMDocumentProcessor:
    """中医文档处理器"""
    
    def __init__(self):
        # 中医专业词汇（jieba分词优化）
        self.tcm_terms = [
            # 病症
            "感冒", "高血压", "糖尿病", "咳嗽", "月经失调", "腰痛", "失眠", "湿疹", "消化性溃疡",
            "头痛", "发热", "便秘", "腹泻", "胸闷", "心悸", "眩晕", "乏力", "食欲不振",
            
            # 中医理论
            "气血", "阴阳", "五脏六腑", "经络", "穴位", "脏腑", "精气神", "卫气营血", "三焦",
            "肝气郁结", "脾胃虚弱", "肾阳不足", "心火上炎", "肺热咳嗽", "湿热下注",
            
            # 诊断方法
            "望闻问切", "四诊", "脉诊", "舌诊", "问诊", "切诊", "八纲辨证", "脏腑辨证",
            
            # 治疗方法
            "针灸", "推拿", "拔罐", "刮痧", "艾灸", "中药", "汤剂", "丸散膏丹", "食疗",
            
            # 常用方剂
            "麻黄汤", "桂枝汤", "小柴胡汤", "四物汤", "六味地黄丸", "逍遥散", "补中益气汤",
            "清热解毒", "活血化瘀", "益气养血", "滋阴润燥", "温阳化气", "疏肝理气"
        ]
        
        # 添加专业词汇到jieba词典
        for term in self.tcm_terms:
            jieba.add_word(term)
            
        # 中医文档结构关键词
        self.section_keywords = {
            'symptoms': ['症状', '主症', '症见', '临床表现', '主要症状', '症候'],
            'etiology': ['病因', '病机', '发病机制', '致病因素', '病理'],
            'diagnosis': ['诊断', '辨证', '证候', '证型', '诊断要点'],
            'treatment': ['治疗', '治法', '方药', '处方', '用药', '方剂'],
            'prescription': ['方剂', '处方', '药物组成', '君臣佐使', '配伍'],
            'nursing': ['护理', '调护', '注意事项', '预防', '保健']
        }
        
        # 无用信息过滤模式
        self.noise_patterns = [
            r'第\d+页',  # 页码
            r'共\d+页',  # 总页数  
            r'\d{4}年\d{1,2}月\d{1,2}日',  # 日期
            r'作者：.*',  # 作者信息
            r'来源：.*',  # 来源信息
            r'参考文献.*',  # 参考文献
            r'目\s*录',  # 目录
            r'^[\s\d\.\(\)]*$',  # 纯数字行
            r'^\s*$'  # 空行
        ]
        
    def extract_text_from_docx(self, docx_path: str) -> str:
        """从docx文件提取纯文本"""
        try:
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
                    
            # 处理表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"提取docx文本失败 {docx_path}: {e}")
            return ""
    
    def clean_medical_text(self, text: str) -> str:
        """清洗中医文本"""
        if not text:
            return ""
            
        # 1. 去除噪音信息
        cleaned_text = text
        for pattern in self.noise_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.MULTILINE)
        
        # 2. 标准化标点符号
        cleaned_text = re.sub(r'[。，；：！？]', '。', cleaned_text)
        cleaned_text = re.sub(r'[（）\(\)]', '', cleaned_text)
        cleaned_text = re.sub(r'[\[\]【】]', '', cleaned_text)
        
        # 3. 去除多余空白
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        cleaned_text = re.sub(r'\n\s*\n', '\n', cleaned_text)
        
        # 4. 统一中医术语
        # 症状统一
        cleaned_text = re.sub(r'头疼|脑袋疼', '头痛', cleaned_text)
        cleaned_text = re.sub(r'发烧|身热', '发热', cleaned_text)
        cleaned_text = re.sub(r'拉肚子|泄泻', '腹泻', cleaned_text)
        cleaned_text = re.sub(r'睡不着|不寐', '失眠', cleaned_text)
        
        # 5. 去除HTML标签（如果有）
        cleaned_text = re.sub(r'<[^>]+>', '', cleaned_text)
        
        return cleaned_text.strip()
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """提取文档的不同部分（症状、病因、治疗等）"""
        sections = {
            'symptoms': '',
            'etiology': '', 
            'diagnosis': '',
            'treatment': '',
            'prescription': '',
            'nursing': '',
            'other': ''
        }
        
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 判断是否为节标题
            section_found = False
            for section_type, keywords in self.section_keywords.items():
                for keyword in keywords:
                    if keyword in line and len(line) < 20:  # 标题通常较短
                        current_section = section_type
                        section_found = True
                        break
                if section_found:
                    break
            
            if not section_found:
                sections[current_section] += line + '\n'
        
        # 清理空的部分
        for key in sections:
            sections[key] = sections[key].strip()
            
        return sections
    
    def chunk_tcm_document(self, text: str, chunk_size: int = 400, overlap: int = 100) -> List[Dict[str, Any]]:
        """中医文档智能分块"""
        if not text:
            return []
            
        # 1. 按结构分块
        sections = self.extract_sections(text)
        chunks = []
        
        # 2. 对每个部分进行分块
        for section_type, section_text in sections.items():
            if not section_text:
                continue
                
            section_chunks = self._chunk_by_sentences(
                section_text, 
                chunk_size=chunk_size, 
                overlap=overlap,
                section_type=section_type
            )
            chunks.extend(section_chunks)
        
        # 3. 如果结构分块失败，使用基础分块
        if not chunks:
            chunks = self._chunk_by_sentences(text, chunk_size, overlap, 'general')
        
        return chunks
    
    def _chunk_by_sentences(self, text: str, chunk_size: int, overlap: int, section_type: str) -> List[Dict[str, Any]]:
        """按句子分块"""
        if not text:
            return []
            
        # 按句号分割
        sentences = [s.strip() + '。' for s in text.split('。') if s.strip()]
        if not sentences:
            return []
        
        chunks = []
        current_chunk = ""
        
        i = 0
        while i < len(sentences):
            # 添加句子到当前块
            sentence = sentences[i]
            
            # 检查是否超过大小限制
            if len(current_chunk + sentence) > chunk_size and current_chunk:
                # 保存当前块
                chunks.append({
                    'text': current_chunk.strip(),
                    'section_type': section_type,
                    'chunk_index': len(chunks),
                    'char_count': len(current_chunk)
                })
                
                # 计算重叠部分
                overlap_text = self._get_overlap_text(current_chunk, overlap)
                current_chunk = overlap_text + sentence
            else:
                current_chunk += sentence
            
            i += 1
        
        # 添加最后一个块
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'section_type': section_type,
                'chunk_index': len(chunks),
                'char_count': len(current_chunk)
            })
        
        return chunks
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """获取重叠文本"""
        if len(text) <= overlap_size:
            return text
        return text[-overlap_size:]
    
    def process_tcm_document(self, docx_path: str, chunk_size: int = 400, overlap: int = 100) -> Dict[str, Any]:
        """处理单个中医文档的完整流程"""
        result = {
            'success': False,
            'filename': os.path.basename(docx_path),
            'raw_text': '',
            'cleaned_text': '',
            'sections': {},
            'chunks': [],
            'stats': {},
            'error_message': ''
        }
        
        try:
            # 1. 提取文本
            raw_text = self.extract_text_from_docx(docx_path)
            if not raw_text:
                result['error_message'] = "无法提取文档内容"
                return result
            
            result['raw_text'] = raw_text
            
            # 2. 清洗文本
            cleaned_text = self.clean_medical_text(raw_text)
            result['cleaned_text'] = cleaned_text
            
            # 3. 提取结构化部分
            sections = self.extract_sections(cleaned_text)
            result['sections'] = sections
            
            # 4. 生成分块
            chunks = self.chunk_tcm_document(cleaned_text, chunk_size, overlap)
            result['chunks'] = chunks
            
            # 5. 统计信息
            result['stats'] = {
                'raw_char_count': len(raw_text),
                'cleaned_char_count': len(cleaned_text),
                'sections_count': len([s for s in sections.values() if s]),
                'chunks_count': len(chunks),
                'avg_chunk_size': sum(c['char_count'] for c in chunks) / len(chunks) if chunks else 0
            }
            
            result['success'] = True
            logger.info(f"文档处理成功: {result['filename']}, 生成{len(chunks)}个分块")
            
        except Exception as e:
            result['error_message'] = str(e)
            logger.error(f"文档处理失败 {docx_path}: {e}")
        
        return result
    
    def batch_process_documents(self, docx_files: List[str]) -> List[Dict[str, Any]]:
        """批量处理文档"""
        results = []
        
        for docx_file in docx_files:
            logger.info(f"处理文档: {docx_file}")
            result = self.process_tcm_document(docx_file)
            results.append(result)
            
            if result['success']:
                print(f"✅ {result['filename']}: {result['stats']['chunks_count']}个分块")
            else:
                print(f"❌ {result['filename']}: {result['error_message']}")
        
        return results

# 使用示例
if __name__ == "__main__":
    processor = TCMDocumentProcessor()
    
    # 测试单个文档
    test_file = "/opt/tcm-ai/all_tcm_docs/感冒.docx"
    if os.path.exists(test_file):
        result = processor.process_tcm_document(test_file)
        if result['success']:
            print(f"处理成功: {result['stats']}")
            print(f"前3个分块:")
            for i, chunk in enumerate(result['chunks'][:3]):
                print(f"分块{i+1}: {chunk['text'][:100]}...")
        else:
            print(f"处理失败: {result['error_message']}")